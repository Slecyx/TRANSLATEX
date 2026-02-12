from docx import Document
from .translator import translate_text

def translate_docx(input_path: str, output_path: str, target_lang: str, progress_callback=None):
    """Translates a DOCX file to the target language with batch processing."""
    from .translator import translate_text, batch_translate_texts
    doc = Document(input_path)
    
    # Collect all translatable headers
    # Mapping mechanism: list of (object, text)
    translatables = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            translatables.append(para)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        translatables.append(para)
    
    total_items = len(translatables)
    if total_items == 0:
        doc.save(output_path)
        return output_path

    # Process in batches
    BATCH_SIZE = 20
    for i in range(0, total_items, BATCH_SIZE):
        batch = translatables[i:i+BATCH_SIZE]
        texts = [item.text for item in batch]
        
        translated_texts = batch_translate_texts(texts, target_lang)
        
        for item, new_text in zip(batch, translated_texts):
            item.text = new_text
            
        # Update progress
        current_progress = min(i + BATCH_SIZE, total_items)
        if progress_callback:
            progress_callback(current_progress, total_items)

    doc.save(output_path)
    return output_path
