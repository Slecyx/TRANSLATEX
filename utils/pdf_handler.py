from pdf2docx import Converter
from docx import Document
import subprocess
import os
from .translator import translate_text

def translate_pdf(input_path: str, output_path: str, target_lang: str, progress_callback=None, convert_to_pdf=True):
    """
    Translates a PDF file to the target language via DOCX conversion.
    Returns the path to the final output file (PDF if possible, else DOCX).
    """
    # ... (previous code specific to DOCX conversion stays same) ...
    # Step 1: Convert PDF to DOCX
    docx_path = input_path.replace('.pdf', '.docx')
    if os.path.exists(docx_path):
        os.remove(docx_path)
        
    cv = Converter(input_path)
    cv.convert(docx_path)
    cv.close()
    
    # Step 2: Translate DOCX content
    from .translator import batch_translate_texts, translate_text
    doc = Document(docx_path)
    
    translatables = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            translatables.append(para)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        translatables.append(para)
                        
    total_items = len(translatables)
    if total_items > 0:
        BATCH_SIZE = 20
        for i in range(0, total_items, BATCH_SIZE):
            batch = translatables[i:i+BATCH_SIZE]
            texts = [item.text for item in batch]
            
            translated_texts = batch_translate_texts(texts, target_lang)
            
            for item, new_text in zip(batch, translated_texts):
                item.text = new_text
                
            current_progress = min(i + BATCH_SIZE, total_items)
            if progress_callback:
                progress_callback(current_progress, total_items)
                
    translated_docx_path = output_path.replace('.pdf', '_translated.docx')
    doc.save(translated_docx_path)
    
    if not convert_to_pdf:
        return translated_docx_path

    # Step 3: Convert back to PDF using LibreOffice (if available)
    try:
        # Check if libreoffice is available
        subprocess.run(['libreoffice', '--version'], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        # Convert to PDF
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', translated_docx_path, '--outdir', os.path.dirname(output_path)], check=True)
        
        # The output file will have same name as translated_docx_path but with .pdf estension
        expected_pdf_path = translated_docx_path.replace('.docx', '.pdf')
        
        if os.path.exists(expected_pdf_path):
             # Rename to desired output path
            if os.path.exists(output_path):
                os.remove(output_path)
            os.rename(expected_pdf_path, output_path)
            return output_path
        else:
             print("LibreOffice conversion failed silently, returning DOCX.")
             return translated_docx_path

    except (subprocess.CalledProcessError, FileNotFoundError):
        print("LibreOffice not found or failed. Returning translated DOCX.")
        return translated_docx_path
